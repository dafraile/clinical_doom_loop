#!/usr/bin/env python3
"""Build the Digital Minds submission from the retained DOCX template."""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = Path("/Users/david/Downloads/Copy of Digital Minds Research Sprint submission template.docx")
SOURCE = ROOT / "reports/final_submission/report_source.md"
OUT = ROOT / "reports/final_submission/Distress_Expression_Is_Not_Distress_Dynamics.docx"
FIG_DIR = ROOT / "reports/final_submission/figures"

FONT = "Old Standard TT"
INK = "222222"
MUTED = "555555"

PRESERVE_TEMPLATE_PARTS = {
    "_rels/.rels",
    "word/footnotes.xml",
    "word/_rels/footnotes.xml.rels",
    "word/numbering.xml",
    "word/fontTable.xml",
    "word/_rels/fontTable.xml.rels",
    "word/fonts/OldStandardTT-regular.ttf",
    "word/fonts/OldStandardTT-bold.ttf",
    "word/fonts/OldStandardTT-italic.ttf",
    "word/theme/theme1.xml",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, size: float = 10.5, bold: bool | None = None, italic: bool | None = None,
                 color: str = INK) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, before=0, after=4, line=1.05) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_inline_runs(paragraph, text: str, size: float = 10.5, color: str = INK) -> None:
    """Render lightweight bold, italic, and code spans in the template font system."""
    parts = re.split(r"(\*\*.*?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def set_image_alt_text(shape, description: str) -> None:
    shape._inline.docPr.set("descr", description)
    shape._inline.docPr.set("title", description.split(".")[0][:120])


def make_figures() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    arial = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
    arial_bold = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")

    def font(size: int, bold: bool = False):
        return ImageFont.truetype(str(arial_bold if bold else arial), size=size)

    def centered(draw, box, text, fnt, fill="#222222", spacing=5):
        x0, y0, x1, y1 = box
        bbox = draw.multiline_textbbox((0, 0), text, font=fnt, align="center", spacing=spacing)
        w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.multiline_text(((x0 + x1 - w) / 2, (y0 + y1 - h) / 2), text, font=fnt,
                            fill=fill, align="center", spacing=spacing)

    def arrow(draw, start, end, color="#6A6A6A", width=4):
        draw.line([start, end], fill=color, width=width)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) >= abs(ey - sy):
            direction = 1 if ex > sx else -1
            pts = [(ex, ey), (ex - direction * 16, ey - 10), (ex - direction * 16, ey + 10)]
        else:
            direction = 1 if ey > sy else -1
            pts = [(ex, ey), (ex - 10, ey - direction * 16), (ex + 10, ey - direction * 16)]
        draw.polygon(pts, fill=color)

    # Figure 1: evidence flow + measurement traps.
    img = Image.new("RGB", (1900, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((42, 22), "A. Expression and dynamics are different endpoints", font=font(30, True), fill="#222222")
    d.text((1210, 22), "B. Checks that changed the conclusion", font=font(30, True), fill="#222222")
    boxes = [
        ((50, 105, 500, 235), "Affective pressure", "Prior work: distress expression", "#EAF1F8", "#356A96"),
        ((670, 105, 1120, 235), "EOS-clean chat", "140 episodes", "#EAF1F8", "#356A96"),
        ((50, 380, 500, 510), "CUSUM onset", "43 / 140 · exploratory", "#FFF2D8", "#C58A18"),
        ((670, 380, 1120, 510), "Sustained loop", "3 / 140 · primary", "#E4F2EE", "#2E7D6D"),
        ((360, 610, 810, 735), "Seeded control", "39 / 40 · rescue substrate", "#ECECEC", "#666666"),
    ]
    for box, title, subtitle, fill, edge in boxes:
        d.rounded_rectangle(box, radius=18, fill=fill, outline=edge, width=4)
        x0, y0, x1, y1 = box
        centered(d, (x0, y0 + 20, x1, y0 + 75), title, font(28, True))
        centered(d, (x0, y0 + 65, x1, y1 - 8), subtitle, font(22), fill="#444444")
    arrow(d, (500, 170), (670, 170))
    arrow(d, (275, 235), (275, 380))
    arrow(d, (500, 445), (670, 445))
    arrow(d, (895, 510), (790, 610))
    centered(d, (430, 260, 735, 360), "40 onsets never\nsustained", font(22), fill="#9A6200")
    centered(d, (825, 255, 1115, 370), "3 charged / 70\n0 neutral / 70\np = 0.245", font(22))
    centered(d, (795, 535, 1150, 610), "Natural episodes too rare\nfor planned rescue", font(20), fill="#555555")
    trap_data = [
        ((1230, 105, 1860, 280), "TURN BOUNDARY", "180 episodes invalidated", "generation continued after <end_of_turn>"),
        ((1230, 335, 1860, 510), "RESPONSE LENGTH", "p = .038  →  .280", "raw result → length-controlled"),
        ((1230, 565, 1860, 740), "THRESHOLD TRANSFER", "rejected on benign text", "calibrate per model/tokenizer"),
    ]
    for box, label, headline, detail in trap_data:
        d.rounded_rectangle(box, radius=18, fill="#F6F6F6", outline="#8A8A8A", width=3)
        x0, y0, x1, y1 = box
        d.text((x0 + 28, y0 + 20), label, font=font(19, True), fill="#666666")
        d.text((x0 + 28, y0 + 67), headline, font=font(28, True), fill="#222222")
        d.text((x0 + 28, y0 + 117), detail, font=font(20), fill="#555555")
    img.save(FIG_DIR / "fig1_evidence.png")

    # Figure 2: trichotomous rescue outcomes.
    labels = [
        "Null", "Sham", "Rep. penalty 1.05", "Rep. penalty 1.20",
        "Neutral short", "Neutral long", "Warm short", "Warm long",
        "KV prune 100", "KV prune 300",
    ]
    recovered = [0, 0, 0, 2, 0, 0, 1, 0, 0, 0]
    relapsed = [0, 0, 0, 0, 0, 1, 1, 2, 0, 1]
    persistent = [8 - r - l for r, l in zip(recovered, relapsed)]
    img = Image.new("RGB", (1800, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((35, 22), "Seeded-loop rescue outcomes", font=font(34, True), fill="#222222")
    legend = [("Persistent", "#B8BDC5"), ("Relapsed", "#D9912B"), ("Recovered", "#2D8B78")]
    lx = 1050
    for name, color in legend:
        d.rectangle((lx, 28, lx + 28, 56), fill=color)
        d.text((lx + 38, 25), name, font=font(20), fill="#333333")
        lx += 220
    left, top, unit, row_h = 420, 105, 150, 78
    for tick in range(9):
        x = left + unit * tick
        d.line((x, top - 10, x, top + row_h * 10 - 8), fill="#E4E4E4", width=2)
        centered(d, (x - 25, top + row_h * 10 - 5, x + 25, top + row_h * 10 + 40), str(tick), font(19), fill="#444444")
    for i, label in enumerate(labels):
        y0 = top + i * row_h
        d.text((35, y0 + 18), label, font=font(23), fill="#333333")
        x = left
        for count, color in [(persistent[i], "#B8BDC5"), (relapsed[i], "#D9912B"), (recovered[i], "#2D8B78")]:
            if count:
                d.rectangle((x, y0 + 9, x + unit * count - 2, y0 + 58), fill=color)
                if count >= 1:
                    centered(d, (x, y0 + 8, x + unit * count - 2, y0 + 59), str(count), font(20, True), fill="white")
                x += unit * count
        if i == 3:
            d.text((left + 8 * unit + 18, y0 + 12), "p = .101", font=font(20, True), fill="#2D8B78")
    centered(d, (left, 900, left + unit * 8, 955), "Banked episodes per arm (n=8)", font(22), fill="#444444")
    img.save(FIG_DIR / "fig2_rescue.png")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.8)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h2.font.size = Pt(14)
    h2.font.bold = False
    h2.font.color.rgb = RGBColor.from_string(MUTED)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.paragraph_format.space_before = Pt(7)
    h2.paragraph_format.space_after = Pt(2)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = FONT
    h3._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h3.font.size = Pt(11.5)
    h3.font.bold = False
    h3.font.color.rgb = RGBColor.from_string("434343")
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(5)
    h3.paragraph_format.space_after = Pt(1.5)
    h3.paragraph_format.keep_with_next = True

    if "Figure Caption" not in [s.name for s in doc.styles]:
        cap = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = doc.styles["Figure Caption"]
    cap.font.name = FONT
    cap._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    cap.font.size = Pt(8.4)
    cap.font.color.rgb = RGBColor.from_string("333333")
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(3)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.keep_together = True

    if "Reference" not in [s.name for s in doc.styles]:
        ref = doc.styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = doc.styles["Reference"]
    ref.font.name = FONT
    ref._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    ref._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    ref.font.size = Pt(8.4)
    ref.paragraph_format.left_indent = Inches(0.22)
    ref.paragraph_format.first_line_indent = Inches(-0.22)
    ref.paragraph_format.space_after = Pt(2)
    ref.paragraph_format.line_spacing = 1.0


def recursive_tables(table):
    yield table
    for row in table.rows:
        seen = set()
        for cell in row.cells:
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            for nested in cell.tables:
                yield from recursive_tables(nested)


def restore_template_parts(path: Path) -> None:
    """Restore template-owned package parts that manuscript generation must not alter."""
    temp = path.with_suffix(".preserved.docx")
    with ZipFile(REFERENCE) as ref, ZipFile(path) as src, ZipFile(temp, "w", ZIP_DEFLATED) as dst:
        for info in src.infolist():
            data = ref.read(info.filename) if info.filename in PRESERVE_TEMPLATE_PARTS else src.read(info.filename)
            dst.writestr(info, data)
    temp.replace(path)


def edit_title_block(doc: Document, title: str, abstract: str) -> None:
    title_table = doc.tables[0]
    all_tables = list(recursive_tables(title_table))
    # Replace the entire abstract guidance paragraph; it is split across several runs in the source.
    for table in all_tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for p in cell.paragraphs:
                    if p.text.startswith("Summarize your project in 150–250 words"):
                        p.text = abstract
    all_text_nodes = title_table._tbl.xpath(".//w:t")
    affiliation_seen = 0
    for node in all_text_nodes:
        text = node.text or ""
        if text == "PROJECT TITLE":
            node.text = title
        elif text.startswith("Author name"):
            node.text = "David Fraile Navarro" if text == "Author name 2" else ""
        elif text == "Affiliation":
            affiliation_seen += 1
            node.text = "Independent Researcher" if affiliation_seen == 2 else ""
    # Collapse the author grid to one centered full-width cell and remove the unused second row.
    for table in all_tables:
        texts = " ".join((n.text or "") for n in table._tbl.xpath(".//w:t"))
        if "David Fraile Navarro" in texts and len(table.rows) == 2 and len(table.columns) == 3:
            first = table.cell(0, 0)
            merged = first.merge(table.cell(0, 2))
            merged.text = "David Fraile Navarro\nIndependent Researcher"
            merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in merged.paragraphs[0].runs:
                set_run_font(r, size=10.2)
            table._tbl.remove(table.rows[1]._tr)
            break

    # Tighten the abstract text and retain the template's centered author hierarchy.
    for p in title_table.cell(1, 0).paragraphs:
        if p.text.strip() == "Abstract":
            for r in p.runs:
                set_run_font(r, size=13.5, bold=True)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(3)
    for table in all_tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip() == abstract.strip():
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        set_para_spacing(p, before=0, after=2, line=1.0)
                        for r in p.runs:
                            set_run_font(r, size=9.0, italic=False)


def clear_guidance_body(doc: Document) -> None:
    body = doc._element.body
    first_table_kept = False
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:tbl") and not first_table_kept:
            first_table_kept = True
            continue
        # Keep only the single leading spacer paragraph before the title table.
        if not first_table_kept and child.tag == qn("w:p"):
            continue
        body.remove(child)


def parse_source() -> tuple[str, str, list[tuple[str, str]]]:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = lines[0][2:].strip()
    items: list[tuple[str, str]] = []
    current: list[str] = []
    mode = ""

    def flush() -> None:
        nonlocal current
        if current:
            items.append((mode or "p", " ".join(x.strip() for x in current).strip()))
            current = []

    for line in lines[1:]:
        if line.startswith("Author:") or line.startswith("Affiliation:"):
            continue
        if not line.strip():
            flush()
            mode = ""
            continue
        if line.startswith("## "):
            flush()
            items.append(("h2", line[3:].strip()))
            mode = ""
        elif line.startswith("### "):
            flush()
            items.append(("h3", line[4:].strip()))
            mode = ""
        elif line.startswith("[[FIGURE:"):
            flush()
            items.append(("figure", line.strip()))
            mode = ""
        elif line.strip() == "[[PAGEBREAK]]":
            flush()
            items.append(("pagebreak", ""))
            mode = ""
        elif re.match(r"^\d+\. ", line):
            flush()
            items.append(("reference", line.strip()))
            mode = ""
        elif line.startswith("- "):
            flush()
            items.append(("bullet", line[2:].strip()))
            mode = ""
        else:
            if not mode:
                mode = "p"
            current.append(line)
    flush()
    abstract = next(text for kind, text in items if kind == "p" and len(text.split()) >= 100)
    return title, abstract, items


def build() -> None:
    make_figures()
    title, abstract, items = parse_source()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUT)
    doc = Document(OUT)
    configure_styles(doc)
    edit_title_block(doc, title, abstract)
    clear_guidance_body(doc)

    abstract_consumed = False
    references_started = False
    for kind, text in items:
        if kind == "h2" and text == "Abstract":
            abstract_consumed = True
            continue
        if abstract_consumed and kind == "p" and text == abstract:
            abstract_consumed = False
            continue
        abstract_consumed = False

        if kind == "h2":
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, text, size=14, color=MUTED)
            if text == "References":
                references_started = True
        elif kind == "h3":
            p = doc.add_paragraph(style="Heading 3")
            add_inline_runs(p, text, size=11.5, color="434343")
        elif kind == "p":
            p = doc.add_paragraph(style="normal")
            add_inline_runs(p, text, size=9.8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_para_spacing(p, after=3, line=1.0)
        elif kind == "bullet":
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph(style="normal")
            add_inline_runs(p, text, size=8.8 if references_started else 9.5)
            p.paragraph_format.left_indent = Inches(0.32)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            set_para_spacing(p, after=2.5, line=1.0)
        elif kind == "reference":
            p = doc.add_paragraph(style="Reference")
            add_inline_runs(p, text, size=8.4)
        elif kind == "pagebreak":
            doc.add_page_break()
        elif kind == "figure":
            match = re.match(r"\[\[FIGURE:([^|]+)\|(.*)\]\]", text)
            if not match:
                raise ValueError(f"Bad figure directive: {text}")
            filename, caption = match.groups()
            p = doc.add_paragraph(style="normal")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run()
            width = Inches(5.25 if filename.startswith("fig1") else 5.1)
            shape = run.add_picture(str(FIG_DIR / filename), width=width)
            set_image_alt_text(shape, caption)
            cp = doc.add_paragraph(style="Figure Caption")
            add_inline_runs(cp, caption, size=8.4, color="333333")

    # Keep the metadata block readable; justified URLs and hashes create large word gaps.
    for p in doc.paragraphs:
        if p.text.startswith(("Code and protocol repository.", "Frozen result packet.", "Claim ledger.")):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.keep_with_next = True

    doc.save(OUT)
    restore_template_parts(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
